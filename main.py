
import os
import uuid
import subprocess
import shutil
import time
import requests
import pyodbc
from datetime import datetime
from pymongo import MongoClient
from duckduckgo_search import DDGS
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from graphviz import Source
import openai
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import uvicorn
from typing import Optional
import logging

# === Setup Logging ===
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# === API Key & Storage Paths ===
openai.api_key = os.getenv("OPENAI_API_KEY")
STORAGE_ROOT = r"\\LANSTAIAPP\Documents\Sessions"
VIDEO_DIR = os.path.join(STORAGE_ROOT, "recordings")
PROCESSED_DIR = os.path.join(VIDEO_DIR, "processed")
CHUNK_DIR = os.path.join(VIDEO_DIR, "chunks")
OUTPUT_DOC_DIR = os.path.join(STORAGE_ROOT, "output_docs")

# === Ensure Required Folders Exist ===
for path in [PROCESSED_DIR, CHUNK_DIR, VIDEO_DIR, OUTPUT_DOC_DIR]:
    os.makedirs(path, exist_ok=True)

from urllib.parse import quote_plus
from pymongo import MongoClient

# === MongoDB Setup ===
mongo_user = quote_plus("LanTech")
mongo_password = quote_plus("L@nc^ere@0012")
mongo_host = "192.168.48.201"
mongo_port = "27017"

MONGO_URI = f"mongodb://{mongo_user}:{mongo_password}@{mongo_host}:{mongo_port}/SuperDB?authSource=admin"

mongo_client = MongoClient(MONGO_URI)
db = mongo_client["sample_db"]
collection = db["test"]


# === SQL Server Setup ===
SQL_CONN_STR = (
    "Driver={ODBC Driver 17 for SQL Server};"
    "Server=183.82.108.211;"
    "Database=SampleDB;"
    "UID=Connectly;"
    "PWD=LT@connect25;"
)

# === FastAPI App ===
app = FastAPI(title="Video Processing API")

# === Pydantic Models ===
class VideoProcessRequest(BaseModel):
    meeting_id: str
    user_id: str

# === Database Initialization ===
def initialize_sql():
    try:
        conn = pyodbc.connect(SQL_CONN_STR)
        cursor = conn.cursor()

        cursor.execute("""
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'tbl_Users')
        CREATE TABLE tbl_Users (
            ID INT IDENTITY(1,1) PRIMARY KEY,
            full_name NVARCHAR(100) NOT NULL,
            email NVARCHAR(100) NOT NULL,
            password NVARCHAR(255) NOT NULL,
            phone_number NVARCHAR(20),
            address NVARCHAR(255),
            country NVARCHAR(50),
            Status BIT DEFAULT 1,
            status_Code CHAR(1) DEFAULT 'u',
            country_code NVARCHAR(10),
            languages NVARCHAR(100),
            agreeToTerms BIT DEFAULT 0,
            Created_At DATETIME DEFAULT GETDATE(),
            Updated_At DATETIME NULL
        )""")

        cursor.execute("""
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'tbl_Meetings')
        CREATE TABLE tbl_Meetings (
            ID UNIQUEIDENTIFIER PRIMARY KEY DEFAULT NEWID(),
            Host_ID INT,
            Meeting_Name NVARCHAR(200),
            Meeting_Type NVARCHAR(50) CHECK (Meeting_Type IN ('CalendarMeeting', 'ScheduleMeeting', 'InstantMeeting')),
            Meeting_Link NVARCHAR(500),
            Status NVARCHAR(50) CHECK (Status IN ('active', 'ended', 'scheduled')) DEFAULT 'active',
            Created_At DATETIME DEFAULT GETDATE(),
            Started_At DATETIME,
            Ended_At DATETIME,
            Is_Recording_Enabled BIT DEFAULT 0,
            Waiting_Room_Enabled BIT DEFAULT 0,
            CONSTRAINT FK_Meetings_Users FOREIGN KEY (Host_ID)
                REFERENCES tbl_Users(ID)
                ON DELETE NO ACTION
                ON UPDATE NO ACTION
        )""")

        conn.commit()
        conn.close()
        logger.info("[SQL] Tables verified.")
    except Exception as e:
        logger.error(f"[ERROR] SQL Initialization failed: {e}")
        raise HTTPException(status_code=500, detail=f"SQL Initialization failed: {str(e)}")

# === Video Processing Functions ===
def compress_and_extract(video_path: str, video_id: str):
    compressed = os.path.join(PROCESSED_DIR, f"compressed_{video_id}.mp4")
    audio = os.path.join(PROCESSED_DIR, f"audio_{video_id}.mp3")
    denoised = os.path.join(PROCESSED_DIR, f"denoised_{video_id}.wav")
    try:
        subprocess.run(["ffmpeg", "-y", "-i", video_path, "-c:v", "libx264", "-crf", "35", "-preset", "veryfast", "-c:a", "aac", "-b:a", "64k", compressed], check=True)
        subprocess.run(["ffmpeg", "-y", "-i", compressed, "-vn", "-ar", "16000", "-ac", "1", "-af", "afftdn", denoised], check=True)
        subprocess.run(["ffmpeg", "-y", "-i", denoised, "-ar", "16000", "-ac", "1", "-b:a", "64k", audio], check=True)
        os.remove(denoised)
        return compressed, audio
    except Exception as e:
        logger.error(f"[ERROR] Compression failed: {e}")
        return None, None

def split_audio_chunks(audio_path: str, video_id: str):
    pattern = os.path.join(CHUNK_DIR, f"{video_id}_chunk_%03d.flac")
    subprocess.run(["ffmpeg", "-y", "-i", audio_path, "-f", "segment", "-segment_time", "300", "-ar", "16000", "-ac", "1", "-c:a", "flac", pattern], check=True)
    return sorted([os.path.join(CHUNK_DIR, f) for f in os.listdir(CHUNK_DIR) if f.startswith(video_id)])

def transcribe_chunks(chunk_paths: list):
    full = ""
    for path in chunk_paths:
        for attempt in range(3):
            try:
                if os.path.getsize(path) > 25 * 1024 * 1024:
                    logger.warning(f"[SKIP] Chunk too large: {path}")
                    break
                with open(path, "rb") as f:
                    response = openai.Audio.transcribe(model="whisper-1", file=f, response_format="text")
                full += response + "\n"
                break
            except Exception as e:
                logger.error(f"[ERROR] Transcription failed: {e}")
                time.sleep(5)
        os.remove(path)
    return full

def get_web_contexts(titles: list):
    contexts = {}
    for title in titles:
        try:
            with DDGS() as ddgs:
                results = ddgs.text(title, max_results=1)
                for r in results:
                    url = r.get("href")
                    html = requests.get(url, timeout=10).text
                    soup = BeautifulSoup(html, "html.parser")
                    text = " ".join(p.get_text() for p in soup.find_all("p")[:6])
                    contexts[title] = text[:2000]
                    break
        except Exception as e:
            logger.warning(f"[SKIP] Web context failed for {title}: {e}")
            continue
    return "\n".join(contexts.values())

def summarize_segment(transcript: str, context: str = ""):
    prompt = f"""
You are a senior documentation and technical writing expert. Your task is to convert the following raw transcript segment into a comprehensive, highly accurate, and formal implementation or study guide based on the subject matter discussed.

The final output must:

- Be structured and formatted according to professional standards for enterprise-level training, onboarding, line pictures, and technical enablement.
- Include step-by-step procedures, clearly numbered and logically ordered.
- Provide real-world tools, technologies, configurations, commands, and screenshots/images (placeholders if needed) relevant to the topic.
- Embed technical examples, use cases, CLI/GUI instructions, and expected outputs or screenshots where applicable.
- Cover common pitfalls, troubleshooting tips, and best practices to ensure full practical understanding.
- Use terminology and instructional depth suitable for readers to gain 100% conceptual and hands-on knowledge of the subject.
- The final document should resemble internal documentation used at organizations like SAP, Oracle, Java, Selenium, AI/ML, Data Science, AWS, Microsoft, or Google — clear, comprehensive, and instructional in tone.

- Additionally, ensure that **for every main topic, you provide 5-10 sentence descriptions** that explain key concepts and their real-world applications. For example, for "Oracle Database" or "Generative AI," give a clear explanation, its use cases, and why it is essential for enterprises. Avoid high-level jargon. Make it practical, applicable, and understandable.

---

OBJECTIVE:

Create a detailed, real-world step-by-step implementation or process guide for [INSERT TOPIC/SUBJECT], designed specifically to support the creation of over 100 technical or comprehension questions. The guide must:

- Reflect real-world tools, technologies, workflows, and industry terminology.
- Break down each phase of the implementation or process logically and sequentially.
- Include practical examples, code snippets (if applicable), key decisions, best practices, and commonly used tools at each step.
- Highlight common challenges or misconceptions, and how they’re addressed in real practice.
- Use terminology and structure that would support SMEs or instructional designers in generating high-quality technical questions based on the guide.
- Avoid abstract or overly generic statements — focus on precision, clarity, and applied knowledge.

---

DOCUMENT FORMAT & STRUCTURE RULES:

1. STRUCTURE
- Use numbered sections and sub-sections (e.g., 1, 1.1, 1.2.1)
- No markdown, emojis, or decorative formatting
- Use plain, formal, enterprise-grade language

2. EACH SECTION MUST INCLUDE:
- A *clear title* and *brief purpose statement*
- *Step-by-step technical or procedural instructions*, including:
    - All relevant tools, platforms, or interfaces used (if any)
    - Any paths, commands, actions, configurations, or API calls involved
    - All required inputs, values, parameters, or dependencies
    - A logical sequence of operations, clearly numbered or separated by actionable steps
    - Tips, warnings, and Important Notes, or expected outcomes where necessary
- **5-10 sentence description** of each main topic, explaining what the concept is, its use cases, and real-world applications. This should be clear and concise for technical audiences to understand why the topic is essential and how it fits into practical workflows.

3. VALIDATION

- Describe how to confirm success (e.g., Expected Outputs, System or Health Checks, Technical and Functional Verifications, Visual Indicators, Fallback/Error Conditions indicators)

4. TROUBLESHOOTING (if applicable)

- Clearly list frequent or known issues that may arise during or after the procedure
- Describe the conditions or misconfigurations that typically lead to each issue
- Provide step-by-step corrective actions or configuration changes needed to resolve each problem
- Mention specific file paths, log viewer tools, console commands, or dashboard areas where errors and diagnostics can be found
- Include example error codes or system messages that help in identifying the issue

5. BEST PRACTICES

- You are a senior technical writer. Based on the following transcript or topic, create a BEST PRACTICES section suitable for formal technical documentation, onboarding materials, or enterprise IT guides.
- Efficiency improvements (e.g., time-saving configurations, automation tips)
- Security or compliance tips (e.g., encryption, IAM roles, audit logging)
- Standard operating procedures (SOPs) used in enterprise environments
- Avoided pitfalls and why they should be avoided
- Format the content using bullet points or short sections for clarity and actionability.
- Avoid vague, obvious, or overly general suggestions — focus on real-world, practical insights derived from field experience or best-in-class implementation norms.

6. CONCLUSION
- Summarize what was implemented or discussed
- Confirm expected outcomes and readiness indicators

---

IMPORTANT:
If the input contains any values such as usernames, IP addresses, server names, passwords, port numbers, or similar technical identifiers — replace their actual content with generic XML-style tags, while preserving the sentence structure and purpose. For example:

- Replace any specific IP address with: <ip>
- Replace any actual password or secret with: <password>
- Replace any actual hostname with: <hostname>
- Replace any actual port number with: <port>
- Replace any username with: <username>
- Replace any email with: <email>

Do NOT alter the sentence structure, meaning, or flow — keep the language intact while swapping the actual values with tags
Do not display or retain real values — just show the placeholder tag. Maintain the original meaning and flow of the instructions.
Format the output as clean, professional documentation, suitable for inclusion in implementation guides, SOPs, or training materials.
Highlight any placeholders in a way that makes it easy for the user to identify where to substitute their own values later.

---

Also:
- Cross-check all tools, commands, file paths, service names, APIs, and utilities with reliable, real-world sources (e.g., official vendor documentation, widely accepted best practices).

 1. If something appears ambiguous, incorrect, or outdated, correct it to its current, supported version.
 2. Use only commands, APIs, or tool names that are verifiably valid and relevant to the topic context.
- Consolidate duplicate or fragmented instructions:
 1. If a step or process is repeated across segments, merge them into a single, complete, and accurate version.
 2. Remove redundancy and preserve the most detailed and correct version of each step.
 3. Do NOT include deprecated or unverifiable content:
 4. Exclude outdated commands, legacy references, or tools no longer maintained.
 5. Replace such content with modern equivalents where available.

- Output the final result as a formal technical guide, with:
  1. Clear section headings
  2. Correct and tested commands/scripts
  3. Accurate tool names and workflows
  4. Logical flow suitable for developers, engineers, or IT teams

---

COMBINED INPUT:
\"\"\"{transcript}\n\n{context}\"\"\"

---

FINAL INSTRUCTION:
Return only the fully formatted implementation or process guide includes below

- A clear, descriptive title
- A concise purpose statement or overview
- Prerequisites and tools required
- Numbered step-by-step instructions with:
   1. Commands, paths, configuration settings, or code blocks (as needed)
   2. GUI or CLI actions explained clearly
   3. Expected inputs, parameters, or options
   4. Confirmation of success (outputs, logs, tests, or validation steps)
   5. Troubleshooting (common issues, causes, and resolutions — if applicable)
   6. Best Practices (efficiency, reliability, security — if applicable)
   7. **Include a mind map diagram in DOT format enclosed in triple backticks at the end**
   8. **Insert chart/diagram placeholders inline to represent where the visual mind map image should appear**

- Replace any real usernames, IP addresses, passwords, ports, or hostnames with <username>, <ip>, <password>, <port>, or <hostname> where needed.
- Eliminate all redundant or outdated, abused content. Only use valid and current tools and commands.

End Document with Standardized "Suggested Next Steps" Note  
*Suggested next steps: No specific next steps mentioned in this segment.*
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",  # Updated to valid model (gpt-4.1-nano is not a known OpenAI model)
            messages=[
                {"role": "system", "content": "You are a technical documentation assistant trained to summarize training meetings."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=3000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"[ERROR] Summary generation failed: {e}")
        return "Summary generation failed."

def extract_dot_code(text: str):
    if "```dot" in text:
        start = text.find("```dot") + 7
        end = text.find("```", start)
        return text[start:end].strip()
    return None

def generate_graphviz_image(dot_code: str, path: str):
    try:
        s = Source(dot_code)
        return s.render(filename=path, format="png", cleanup=True) + ".png"
    except Exception as e:
        logger.error(f"[ERROR] DOT render failed: {e}")
        return None

def save_docx(content: str, image_path: Optional[str], filename: str):
    doc = Document()
    doc.add_heading("Summary Document", 0)
    if "```dot" in content:
        content = content.split("```dot")[0].strip()
    for line in content.splitlines():
        doc.add_paragraph(line)
    if image_path and os.path.exists(image_path):
        doc.add_page_break()
        doc.add_heading("Mind Map", level=1)
        doc.add_picture(image_path, width=Inches(6))
    filepath = os.path.join(OUTPUT_DOC_DIR, filename)
    doc.save(filepath)
    return filepath

def save_transcript_docx(transcript: str, filename: str):
    doc = Document()
    doc.add_heading("Transcript", 0)
    doc.add_paragraph(transcript)
    path = os.path.join(OUTPUT_DOC_DIR, filename)
    doc.save(path)
    return path

async def process_video(video_path: str, meeting_id: str, user_id: str):
    original_filename = os.path.splitext(os.path.basename(video_path))[0]
    filename_prefix = f"{meeting_id}_{user_id}_{original_filename}"

    if collection.find_one({"video_path": video_path}):
        logger.info("[SKIP] Already processed.")
        return {"status": "skipped", "message": "Video already processed"}

    compressed, audio = compress_and_extract(video_path, filename_prefix)
    if not compressed or not audio:
        raise HTTPException(status_code=500, detail="Compression failed")

    chunks = split_audio_chunks(audio, filename_prefix)
    transcript = transcribe_chunks(chunks)
    if not transcript.strip():
        logger.info("[SKIP] Empty transcript.")
        raise HTTPException(status_code=400, detail="Empty transcript")

    context = get_web_contexts(["General AI"])
    summary = summarize_segment(transcript, context)
    dot_code = extract_dot_code(summary)
    mindmap = generate_graphviz_image(dot_code, os.path.join(OUTPUT_DOC_DIR, f"{filename_prefix}_mindmap")) if dot_code else None

    doc1 = save_transcript_docx(transcript, f"{filename_prefix}_transcript.docx")
    doc2 = save_docx(summary, mindmap, f"{filename_prefix}_summary.docx")

    collection.insert_one({
        "video_path": os.path.abspath(video_path),
        "meeting_id": meeting_id,
        "user_id": user_id,
        "transcript_doc_path": os.path.abspath(doc1),
        "summary_doc_path": os.path.abspath(doc2),
        "mindmap_image_path": os.path.abspath(mindmap) if mindmap else None,
        "timestamp": datetime.now()
    })

    os.remove(compressed)
    os.remove(audio)
    logger.info("[DONE] Process complete.")
    return {
        "status": "success",
        "transcript_doc": doc1,
        "summary_doc": doc2,
        "mindmap_image": mindmap
    }

# === API Endpoints ===
@app.on_event("startup")
async def startup_event():
    initialize_sql()

@app.post("/upload-video/")
async def upload_video(file: UploadFile = File(...), meeting_id: str = "", user_id: str = ""):
    try:
        # Validate file type
        if not file.filename.endswith(('.mp4', '.mov', '.avi')):
            raise HTTPException(status_code=400, detail="Unsupported file format. Use .mp4, .mov, or .avi")

        # Save uploaded video
        video_id = str(uuid.uuid4())
        video_path = os.path.join(VIDEO_DIR, f"{video_id}.mp4")  # Fixed typo: video_id instead of video hous_id
        with open(video_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Process video
        result = await process_video(video_path, meeting_id, user_id)
        return JSONResponse(content=result)
    except Exception as e:
        logger.error(f"[ERROR] Video upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload/")
async def upload_video_alias(file: UploadFile = File(...), meeting_id: str = "", user_id: str = ""):
    logger.info("Received request to /upload; redirecting to /upload-video/")
    return await upload_video(file, meeting_id, user_id)

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

# === Run the App ===
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
